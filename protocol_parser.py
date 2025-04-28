import argparse
import os
import re
import json
from docx import Document
from docx.enum.text import WD_UNDERLINE

# Regex pattern to match speaker lines, with optional angle brackets and parentheses
speaker_line_pattern = re.compile(r'^(<+)?(.+?)(\s*\(.*?\))?:(>+)?$')
# Regex pattern to remove angle brackets <<...>>
remove_angle_brackets_pattern = re.compile(r'<<.*?>>')
# Regex pattern to match content within parentheses
parentheses_pattern = re.compile(r'\(.*?\)')
# Regex pattern to match any character that is not a Hebrew letter
hebrew_letters_pattern = re.compile(r'[^\u0590-\u05FF]')
# Regex pattern to split sentences at sentence end punctuation
sentence_end_pattern = re.compile(r'(?<=[.!?])\s+(?=[^\s])')
# Regex pattern to tokenize sentences, matching words, numbers, and punctuation
tokenize_pattern = re.compile(r'\w+["\']?\w*|\d+|[\".,;:!?()\[\]{}\-%&@#$*+=/<>~`|\\]')
# Regex pattern to tokenize into parts of digits and letters
number_pattern = re.compile(r'\d+|\D+')
# Regex pattern to match any Latin letter (English alphabet)
alphabet_regex = re.compile(r'[A-Za-z]')
# Regex pattern to match sequences of dashes with optional spaces (e.g., '--', '- -', etc.)
dash_sequence_regex = re.compile(r'(-\s*){2,}')  # Matches sequences of dashes with optional spaces
# Regex pattern to match various types of dashes for normalization
dash_substitution_regex = re.compile(r'[‐‑‒–—―−]')  # Matches various types of dashes for normalization


class Sentence:
    """Represents a sentence spoken by a speaker in the protocol."""

    def __init__(self, speaker_name, sentence_text):
        self.speaker_name = speaker_name
        self.sentence_text = sentence_text


class Protocol:
    """Represents a protocol document, containing sentences."""

    def __init__(self, protocol_name, knesset_number, protocol_type, protocol_number):
        self.protocol_name = protocol_name
        self.knesset_number = knesset_number
        self.protocol_type = protocol_type
        self.protocol_number = protocol_number
        self.sentences = []


def parse_file_name(file_name):
    """
       Parses the file name to extract the Knesset number and protocol type.

       Args:
           file_name (str): The name of the file.

       Returns:
           tuple: A tuple containing the Knesset number (int) and the protocol type (str).
       """
    # Remove extension
    name = os.path.splitext(file_name)[0]
    parts = name.split('_')
    knesset_number = int(parts[0])
    protocol_type_code = parts[1]
    if protocol_type_code == 'ptm':
        protocol_type = 'plenary'
    elif protocol_type_code == 'ptv':
        protocol_type = 'committee'
    else:
        protocol_type = 'unknown'
    return knesset_number, protocol_type


hebrew_numbers = {
    'אחד': 1, 'אחת': 1, 'ראשון': 1, 'שניים': 2, 'שני': 2, 'שתיים': 2, 'שתי': 2, 'שניה': 2,
    'שלושה': 3, 'שלוש': 3, 'שלישי': 3,
    'ארבעה': 4, 'ארבע': 4, 'רביעי': 4,
    'חמישה': 5, 'חמש': 5, 'חמישי': 5,
    'שישה': 6, 'שש': 6, 'שישי': 6,
    'שבעה': 7, 'שבע': 7, 'שביעי': 7,
    'שמונה': 8, 'שמיני': 8,
    'תשעה': 9, 'תשע': 9, 'תשיעי': 9,
    'עשרה': 10, 'עשר': 10, 'עשירי': 10,
    'עשרים': 20, 'שלושים': 30, 'ארבעים': 40, 'חמישים': 50, 'שישים': 60, 'שבעים': 70, 'שמונים': 80, 'תשעים': 90,
    'מאות': 100,
    'מאה': 100, 'מאתיים': 200, 'אלף': 1000
}


def hebrew_words_to_int(words):
    """
        Converts a list of Hebrew number words into an integer.

        Args:
            words (list): List of Hebrew number words.

        Returns:
            int or None: The integer value of the words, or None if conversion fails.
        """
    total = 0
    cur_num = 0
    for word in words:
        base_word = word
        if base_word in hebrew_numbers:
            if base_word == 'מאות':
                total -= cur_num
                total += cur_num * 100
            else:
                cur_num = hebrew_numbers[base_word]
                total += cur_num
        else:
            # Not a number word, stop processing
            break
    return total if total > 0 else None


def extract_protocol_number(doc):
    """
       Extracts the protocol number from the given Word document.

       Args:
           doc (Document): A docx Document object.

       Returns:
           int: The protocol number, or -1 if not found.
       """
    protocol_number = -1
    # Define patterns to search for the protocol number
    patterns = [
        r'<?פרוטוקול\s+(?:מס\'?|מספר)\s*(\d+)>?',  # Match 'פרוטוקול מס' or 'פרוטוקול מספר' followed by digits
        r'הישיבה\s+([א-ת-]+)\s+של'  # Match 'הישיבה <number words> של'
    ]
    for par in doc.paragraphs:
        text = par.text.strip()
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                number_part = match.group(1).strip()

                if not number_part:
                    continue

                # Remove unwanted characters (non-digits, non-Hebrew letters, non-space, non-hyphen)
                number_part = re.sub(r'[^\d\u0590-\u05FF\s\-]', '', number_part)
                # Try to parse as integer
                if re.match(r'^\d+$', number_part):
                    protocol_number = int(number_part)
                    return protocol_number
                else:
                    # Process Hebrew words
                    words = re.split(r'\s+|-', number_part)
                    number_words = []
                    for word in words:
                        word = word.strip()
                        if not word:
                            continue
                        # Remove 'ו' or 'ה' prefixes
                        if word.startswith('ו') or word.startswith('ה'):
                            base_word = word.lstrip('וה')
                        else:
                            base_word = word
                        if base_word in hebrew_numbers:
                            number_words.append(base_word)
                        else:
                            # Stop processing if word not in mapping
                            break
                    # Convert Hebrew words to integer
                    number = hebrew_words_to_int(number_words)
                    if number is not None:
                        protocol_number = number
                        return protocol_number
    return protocol_number


def is_speaker_line(paragraph):
    """
    Checks if a paragraph is a speaker line based on the pattern and if any part of the paragraph is underlined.
    Returns the match object if it's a speaker line, or None otherwise.
    """
    text = paragraph.text.strip()
    # Remove '<< ... >>' tags
    text_cleaned = remove_angle_brackets_pattern.sub('', text).strip()

    # Replace curly quotes with standard quotes
    text_cleaned = text_cleaned.replace("”", "\"").replace("“", "\"").replace("״", "\"")
    text_cleaned = text_cleaned.replace("’", "'").replace("‘", "'").replace("`", "'").replace("´", "'").replace("ʼ",
                                                                                                                "'").replace("‛", "'").replace("׳", "'")
    match = speaker_line_pattern.match(text_cleaned)

    if match and is_paragraph_underlined(paragraph):
        return match, text_cleaned  # Return the match object
    else:
        return None, text_cleaned  # Not a speaker line


def is_paragraph_underlined(paragraph):
    """
    Checks if the paragraph text is underlined, considering direct formatting, styles,
    and base styles.
    """
    # Check if any run is underlined
    for run in paragraph.runs:
        if run.text.strip():
            if is_run_underlined(run):
                return True

    # Check if the paragraph style or its base styles apply underlining
    style = paragraph.style
    while style:
        if style.font.underline in [True, WD_UNDERLINE.SINGLE]:
            return True
        style = style.base_style

    return False


def is_run_underlined(run):
    """
    Returns True if the run is underlined, considering direct formatting and styles.
    """
    # Check direct formatting
    if run.font.underline in [True, WD_UNDERLINE.SINGLE]:
        return True
    # Check character style
    if run.style and run.style.font.underline in [True, WD_UNDERLINE.SINGLE]:
        return True
    return False


def extract_speaker_name(match):
    """
    Extracts the speaker's name using the match object from is_speaker_line.
    """
    text = match.group(2).strip()  # Extract the main content
    # Remove parentheses and their content
    text = parentheses_pattern.sub('', text).strip()
    # Clean speaker name
    name = clean_speaker_name(text)
    return name


def clean_speaker_name(text):
    """
        Cleans and extracts the speaker's name from the text.

        Args:
            text (str): The text containing the speaker's name.

        Returns:
            str: The cleaned speaker name.
        """

    words = text.strip().split()

    # Reverse the list of words for easier processing from the end
    reversed_words = words[::-1]

    # Define a set of valid names starting with "ה"
    valid_names_starting_with_he = {
        'האנה', 'האני', 'הארי', 'הגר', 'הדס', 'הדסה', 'הדר', 'הדרה', 'הוגו', 'הוד', 'הודיה', 'הולי',
        'הורדוס', 'היידי', 'היילי', 'הילאי', 'הילדה', 'הילה', 'הילור', 'הילי', 'הילית', 'הילל', 'הילרי',
        'הינדל', 'הלגה', 'הלל', 'הללי', 'הלן', 'הלנה', 'הלני', 'הני', 'הניה', 'הנרי', 'הנריטה', 'הנרייטה',
        'הענדל', 'העני', 'הקטור', 'הראל', 'הראלה', 'הרברט', 'הרולד', 'הרמיוני', 'הרן', 'הרצל', 'הרשל'
    }

    name_words = []
    stop_words = {'במשרד', 'בממשלה', 'ביטחון', 'בראש',
                  'לביטחון', 'למשטרה', 'לחקלאות', 'פנים', 'לתעשייה', 'והמסחר', 'לסביבה', 'לאוצר', 'לתחבורה',
                  'לתקשורת', 'מודיעין', 'בדרכים',
                  'לתשתיות', 'ללאומיות', 'לעלייה', 'ולקליטה', 'לענייני', 'כלכלה', 'וחברה', 'למנכ"לית', 'למנכ"ל'}

    # Titles to exclude if they appear before the name
    excluded_prefixes = {"ד\"ר", 'ד"ר', "פרופ'", 'עו"ד', 'רב', 'ניצב', 'היו\"ר', 'נצ"מ', 'סא"ל', 'רס"ן', 'תא"ל',
                         'אלוף', 'מר',
                         'גב\'', 'גב"\''}

    for word in reversed_words:
        clean_word = word
        # If the word is a stop word or preposition, we assume the name has ended
        if clean_word in stop_words:
            break
        elif clean_word in excluded_prefixes:
            break  # Skip the prefix
        elif clean_word.startswith('וה'):
            break
        # Check for title-like words (starts with "ה")
        elif len(name_words) >= 2 and clean_word.startswith('ה') and clean_word not in valid_names_starting_with_he:
            break

        name_words.append(word)
        # We can limit the name to a reasonable number of words
        if len(name_words) >= 5:
            break

    # Reverse the name words to get the correct order
    name = ' '.join(name_words[::-1])

    # Remove any trailing colons or punctuation
    name = name.rstrip(':')
    name = name.strip('-')

    return name


def segment_sentences(text):
    """
       Segments the given text into sentences based on sentence-ending punctuation.

       Args:
           text (str): The text to segment.

       Returns:
           list: A list of sentences.
       """
    # Custom sentence segmentation
    # Split at sentence end markers, considering possible Hebrew punctuation
    sentences = sentence_end_pattern.split(text)
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


def is_valid_sentence(sentence):
    """
        Checks if a sentence is valid based on certain criteria.

        Args:
            sentence (str): The sentence to check.

        Returns:
            bool: True if the sentence is valid, False otherwise.
        """
    stripped = hebrew_letters_pattern.sub('', sentence)
    # Check if there are at least some Hebrew letters
    if len(stripped) == 0:
        return False
    # Check if the sentence contains English letters
    if alphabet_regex.search(sentence):
        return False
    # Check for incomplete sentences marked with '---' or similar
    if dash_sequence_regex.search(sentence):
        return False
    return True


def tokenize_sentence(sentence):
    """
       Tokenizes a sentence into words and punctuation marks.

       Args:
           sentence (str): The sentence to tokenize.

       Returns:
           list: A list of tokens.
       """
    # Use regex to split tokens while keeping punctuation and quotes separate
    tokens = tokenize_pattern.findall(sentence)

    # Post-process tokens to separate ending quotes
    processed_tokens = []
    for token in tokens:
        if token.endswith('"') and len(token) > 1:
            # Split the word and the ending quote
            processed_tokens.append(token[:-1])  # Add the word without the quote
            processed_tokens.append('"')  # Add the quote as a separate token
        else:
            processed_tokens.extend(split_mixed_token(token))

    return processed_tokens


def split_mixed_token(token):
    # Check if the token contains both digits and letters
    if any(char.isdigit() for char in token) and any(char.isalpha() for char in token):
        # Split the token into parts of digits and letters
        return number_pattern.findall(token)
    else:
        return [token]


def process_document(file_path, protocol_name, knesset_number, protocol_type):
    """
        Processes a single protocol document and extracts sentences.

        Args:
            file_path (str): Path to the .docx file.
            protocol_name (str): The name of the protocol.
            knesset_number (int): The Knesset number.
            protocol_type (str): The type of protocol ('plenary', 'committee', etc.).

        Returns:
            Protocol: A Protocol object containing extracted sentences.
        """
    document = Document(file_path)
    protocol_number = extract_protocol_number(document)
    protocol = Protocol(protocol_name, knesset_number, protocol_type, protocol_number)
    current_speaker = None
    ignore_speaker = True
    current_speech = []

    paragraphs = document.paragraphs
    for i, par in enumerate(paragraphs):

        match, par_text = is_speaker_line(par)
        if match and (current_speaker is not None or "יו\"ר" in par_text):
            # Save previous speaker's speech
            if current_speaker and current_speech:
                process_speech(protocol, current_speaker, current_speech)
            # Update current_speaker
            speaker_name = extract_speaker_name(match)
            # Determine whether to ignore the speaker
            if 'קריאה' != speaker_name and 'קריאות' != speaker_name:
                ignore_speaker = False
            else:
                ignore_speaker = True
            current_speaker = speaker_name
            current_speech = []
        else:
            text = par_text
            if text and not ignore_speaker:
                current_speech.append(text)
    # Process last speaker's speech
    if current_speaker and current_speech:
        process_speech(protocol, current_speaker, current_speech)
    return protocol


def process_speech(protocol, speaker_name, speech_paragraphs):
    """
       Processes the speech paragraphs of a speaker and adds valid sentences to the protocol.

       Args:
           protocol (Protocol): The Protocol object to which sentences will be added.
           speaker_name (str): The name of the speaker.
           speech_paragraphs (list): List of speech paragraphs as strings.
       """
    try:
        speech_text = ' '.join(speech_paragraphs)
        # Replace all dash-like Unicode characters with a standard hyphen '-'
        speech_text = dash_substitution_regex.sub('-', speech_text)
        sentences = segment_sentences(speech_text)
        for sentence in sentences:
            if is_valid_sentence(sentence):
                tokens = tokenize_sentence(sentence)
                if len(tokens) >= 4:
                    sentence_text = ' '.join(tokens)
                    sentence_obj = Sentence(speaker_name, sentence_text)
                    protocol.sentences.append(sentence_obj)
    except Exception as e:
        print(f"Error processing speech for speaker {speaker_name}: {e}")


def process_all_documents(folder_path, output_file):
    """
       Processes all .docx documents in the specified folder and writes the extracted data to a JSONL file.

       Args:
           folder_path (str): Path to the folder containing .docx files.
           output_file (str): Path to the output JSONL file.
       """
    protocols = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith('.docx'):
            try:
                file_path = os.path.join(folder_path, file_name)
                protocol_name = file_name
                knesset_number, protocol_type = parse_file_name(file_name)
                protocol = process_document(file_path, protocol_name, knesset_number, protocol_type)
                if protocol is not None:
                    protocols.append(protocol)
            except Exception as e:
                print(f"Error processing file {file_name}: {e}")
                continue
    # Write data to JSONL file
    with open(output_file, 'w', encoding='utf-8') as f:
        for protocol in protocols:
            for sentence in protocol.sentences:
                record = {
                    'protocol_name': protocol.protocol_name,
                    'knesset_number': protocol.knesset_number,
                    'protocol_type': protocol.protocol_type,
                    'protocol_number': protocol.protocol_number,
                    'speaker_name': sentence.speaker_name,
                    'sentence_text': sentence.sentence_text
                }
                json_line = json.dumps(record, ensure_ascii=False)
                f.write(json_line + '\n')


def main():
    # Use argparse to parse command-line arguments
    parser = argparse.ArgumentParser(description='Process Knesset protocols.')
    parser.add_argument(
        'input_dir',
        type=str,
        help='Path to the directory containing the input .docx files.'
    )
    parser.add_argument(
        'output_file',
        type=str,
        help='Path to the output JSONL file.'
    )
    args = parser.parse_args()

    # Get the input directory and output file from command-line arguments
    input_dir = args.input_dir
    output_file = args.output_file

    # Ensure the input directory exists
    if not os.path.isdir(input_dir):
        print(f"Error: The directory '{input_dir}' does not exist.")
        return

    # Process all documents
    process_all_documents(input_dir, output_file)
    print(f"Processing completed. Output saved to '{output_file}'.")

